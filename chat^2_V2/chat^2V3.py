import sys, os
import datetime
import json
import time
import google.generativeai as genai
import google.api_core.exceptions

# Obtener el directorio actual (donde se encuentra mi_script.py)
directorio_actual = os.path.dirname(os.path.realpath(__file__))

# Agregar el directorio padre al sys.path
directorio_padre = os.path.abspath(os.path.join(directorio_actual, os.pardir))
sys.path.append(directorio_padre)

# Ahora puedes importar GEMINI_API_KEY desde archivo.py
from archivo import GEMINI_API_KEY


# Configurar la API de generativeai
GEMINI_API_KEY = GEMINI_API_KEY()

def get_last_file_number():
    files = os.listdir("chats_gemini-gemini")
    last_number = len(files)
    return last_number

def iniciar_conversacion(modelo_nombre, historial=None, temperature=1):
    if historial is None:
        historial = []
    genai.configure(api_key=GEMINI_API_KEY)
    generation_config = {
        "temperature": temperature,
        "top_p": 0.95,
        "top_k": 64,
        "max_output_tokens": 8192,
        "response_mime_type": "text/plain",
    }
    safety_settings = [
    {"category": "HARM_CATEGORY_HARASSMENT", "threshold": "BLOCK_NONE",},
    {"category": "HARM_CATEGORY_HATE_SPEECH", "threshold": "BLOCK_NONE",},
    { "category": "HARM_CATEGORY_SEXUALLY_EXPLICIT", "threshold": "BLOCK_NONE",},
    {"category": "HARM_CATEGORY_DANGEROUS_CONTENT", "threshold": "BLOCK_NONE",},
    ]
    model = genai.GenerativeModel(
        model_name=modelo_nombre,
        safety_settings=safety_settings,
        generation_config=generation_config,
    )
    chat_session = model.start_chat(history=historial)
    return chat_session

def enviar_mensaje(chat_session, mensaje):
    max_retries = 5
    for attempt in range(max_retries):
        try:
            response = chat_session.send_message(mensaje)
            return response
        except google.api_core.exceptions.DeadlineExceeded as e:
            print(f"Error: {e}")
            print(f"Reintentando... (Intento {attempt + 1}/{max_retries}, esperando {2 ** attempt} segundos)")
            if attempt < max_retries - 1:
                wait_time = 2 ** attempt
                time.sleep(wait_time)
            else:
                raise e

def guardar_historial(temperature, historial, archivo, modelo1, modelo2):
    if not os.path.exists(archivo):  # Verificar si el archivo ya existe
        # Agregar la configuración inicial al historial
        historial.insert(0, {
            "generation_config": {
                "temperature": temperature,
                "top_p": 0.95,
                "top_k": 64,
                "max_output_tokens": 8192,
                "response_mime_type": "text/plain",
            },
            "safety_settings": [
                {"category": "HARM_CATEGORY_HARASSMENT", "threshold": "BLOCK_NONE",},
                {"category": "HARM_CATEGORY_HATE_SPEECH", "threshold": "BLOCK_NONE",},
                { "category": "HARM_CATEGORY_SEXUALLY_EXPLICIT", "threshold": "BLOCK_NONE",},
                {"category": "HARM_CATEGORY_DANGEROUS_CONTENT", "threshold": "BLOCK_NONE",},
            ],
            "model_names": [modelo1, modelo2]
        })
    with open(archivo, 'w') as f:
        json.dump(historial, f, indent=2)

def cargar_historial(archivo):
    with open(archivo, 'r') as f:
        historial = json.load(f)
    return historial

def main():
    modelo1 = "gemini-1.5-flash"
    modelo2 = "gemini-1.5-pro"
    fecha_hora = datetime.datetime.now().strftime("%Y%m%d")
    historial_archivo = f"chats_gemini-gemini/historial_conversacion_{fecha_hora}_{get_last_file_number()}.json"

    if os.path.exists(historial_archivo):
        historial = cargar_historial(historial_archivo)
    else:
        historial = []

    temperature = float(input("Ingrese la temperatura: "))
    chat_session_1 = iniciar_conversacion(modelo1, historial, temperature=temperature)
    chat_session_2 = iniciar_conversacion(modelo2, historial, temperature=temperature)

    for item in historial:
        if item["message"] and item["response"]:
            print("Modelo 1:", item["message"])
            print("Modelo 2:", item["response"])
    response_2 = None
    while True:
        if response_2 is None:
            # mensaje_1 = input("Modelo 1: ")
            mensaje_1 = """Que opinas del codigo: import json
import unicodedata
import docx

# Leer el archivo JSON

json_file_path = 'chats_gemini-gemini/historial_conversacion_20240601_1.json'

with open(json_file_path, 'r', encoding='utf-8') as file:
    data = json.load(file)

# Función para decodificar caracteres Unicode
def decode_unicode(text):
    return unicodedata.normalize("NFKD", text)

# Crear un nuevo documento Word
doc = docx.Document()
doc.add_heading('Registro de Interacciones', 0)

# Procesar y formatear los datos
for item in data:
    time = decode_unicode(item["time"])
    message = decode_unicode(item["message"])
    response = decode_unicode(item["response"])

    doc.add_heading('Registro', level=1)
    doc.add_paragraph(f"Time: {time}")
    doc.add_paragraph(f"Message: {message}")
    doc.add_heading('Response', level=2)
    doc.add_paragraph(response)

# Guardar el documento
doc_path = f'lectura/historial_conversacion_20240601_1.docx'
doc.save(doc_path)"""
        else:
            mensaje_1 = response_2.text

        response_1 = enviar_mensaje(chat_session_1, mensaje_1)
        historial.append({"time": str(datetime.datetime.now()), "message": mensaje_1, "response": response_1.text})
        print("Modelo 1:", response_1.text)

        time.sleep(30)

        mensaje_2 = response_1.text
        response_2 = enviar_mensaje(chat_session_2, mensaje_2)
        historial.append({"time": str(datetime.datetime.now()), "message": mensaje_2, "response": response_2.text})

        guardar_historial(temperature, historial, historial_archivo, modelo1, modelo2)

        print("Modelo 2:", response_2.text)
        
        time.sleep(30)

if __name__ == "__main__":
    main()
