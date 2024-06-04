import json
import unicodedata
import docx

# Leer el archivo JSON

json_file_path = 'chats_gemini-gemini/historial_conversacion_20240604_14.json'

with open(json_file_path, 'r', encoding='utf-8') as file:
    data = json.load(file)

# Función para decodificar caracteres Unicode
def decode_unicode(text):
    return unicodedata.normalize("NFKD", text)

# Crear un nuevo documento Word
doc = docx.Document()
doc.add_heading('Registro de Interacciones', 0)

# Procesar y formatear los datos
for item in data:
    try:
        time = decode_unicode(item["time"])
        message = decode_unicode(item["message"])
        response = decode_unicode(item["response"])

        doc.add_heading('Registro', level=1)
        doc.add_paragraph(f"Time: {time}")
        doc.add_paragraph(f"Message: {message}")
        doc.add_heading('Response', level=2)
        doc.add_paragraph(response)
    except:
        print("Error en el registro")
        print(item)
# Guardar el documento
doc_path = f'lectura/historial_conversacion_20240604_14.docx'
doc.save(doc_path)