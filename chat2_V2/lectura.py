import json
import unicodedata
import docx

# JSON file name
name = "conversation_history_20240608_23"

# Read the JSON file
json_file_path = f'chats_gemini-gemini/{name}.json'

with open(json_file_path, 'r', encoding='utf-8') as file:
    data = json.load(file)

# Function to decode Unicode characters
def decode_unicode(text):
    return unicodedata.normalize("NFKD", text)

# Create a new Word document
doc = docx.Document()
doc.add_heading('Interaction Log', 0)

# Process and format the data
for item in data:
    try:
        time = decode_unicode(item["time"])
        message = decode_unicode(item["message"])
        response = decode_unicode(item["response"])

        doc.add_heading('Log Entry', level=1)
        doc.add_paragraph(f"Time: {time}")
        doc.add_paragraph(f"Message: {message}")
        doc.add_heading('Response', level=2)
        doc.add_paragraph(response)
    except:
        print("Error in log entry")
        print(item)

# Save the document
doc_path = f'output/{name}.docx'
doc.save(doc_path)

print(f"\nDocument saved at {doc_path}")
